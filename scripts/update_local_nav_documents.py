#!/usr/bin/env python3
"""Migrate the two active internship Word documents to the mapless local-nav scope."""
from __future__ import annotations

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
PLAN = ROOT / "新版实习计划VLN" / "实习计划书1 -徐锦彦.docx"
PATH = ROOT / "新版实习计划VLN" / "GM461-E1文字目标导航与静态避障_一个月完整技术路径.docx"


def cell(table, row: int, column: int, value: str) -> None:
    table.rows[row].cells[column].text = value


def replace_everywhere(document: Document, replacements: dict[str, str]) -> None:
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for item in row.cells:
                paragraphs.extend(item.paragraphs)
    for section in document.sections:
        paragraphs.extend(section.header.paragraphs)
        paragraphs.extend(section.footer.paragraphs)
    for paragraph in paragraphs:
        text = paragraph.text
        updated = text
        for old, new in replacements.items():
            updated = updated.replace(old, new)
        if updated != text:
            paragraph.text = updated


def update_plan() -> None:
    document = Document(PLAN)
    p = document.paragraphs
    p[0].text = "目标可见条件下的短距离局部导航与静态避障"
    p[7].text = (
        "本次实习计划完成一个目标可见条件下的短距离局部导航与静态避障 Demo。使用屈膝升降式移动双臂/四臂机器人 C6-M 和 GM461-E1 RGB-D 相机。输入“寻找椅子”或“找一下背包”后，程序从配置词表解析目标；只有目标已在当前视野中连续三帧确认，才结合深度图计算其相对位置，并在以机器人为中心、持续滚动更新的局部代价地图内规划短距离接近路径。系统不建立、不加载也不依赖预建全局地图，不访问全局观察点，不执行跨房间搜索；当前视野无目标、局部路径被阻或安全状态未知时直接停止并报告失败。最低要求是在录制 RGB-D 或可控模拟数据上完成可重复闭环，真实运动仅在里程计、局部 TF、平台控制接口和全部安全条件获批后低速进行。"
    )
    p[12].text = (
        "先熟悉机器人的开关机、急停、触边、抱闸、人工接管和低速运行方法。使用 Percipio Viewer 和图漾官方 SDK/ROS 示例逐项确认 GM461-E1 的彩色图、深度图、内参、点云、时间戳和 RGB-D 对齐。机器人侧重点确认 odom、base_link、camera_link 的短时坐标链、实时激光、安全状态，以及平台已有局部控制器或厂商 RCS 是否支持局部目标提交、状态查询和停止；不把全局地图或 AMCL 作为项目依赖。"
    )
    p[20].text = (
        "检测到物体后，取检测框中央区域的有效深度中位数并记录质量，再利用真实相机内参反投影为 camera_link 三维点。按图像时间戳查询 camera_link → base_link → odom 的短时 TF，得到局部目标坐标。RGB-D 不同步、内参缺失、深度不足、TF 过期或里程计不可用时不生成局部目标。"
    )
    p[22].text = "4、可见目标确认与滚动局部规划"
    p[24].text = (
        "当前画面只做有限帧检测；目标不在视野中时任务失败，不让机器人为找目标自主前进，也不访问预设观察点。只有同一候选连续三帧在类别、置信度、深度和局部位置上稳定才确认。确认后按目标相对位置计算约 1 m 的停靠点，在 8 m × 8 m 左右的滚动局部代价地图中生成短路径。每次只执行约 0.3—0.5 m，再重新检测目标、更新障碍并重规划；窗口内无安全路径时停止。人员以及 detect_only/disabled 类别不得产生运动任务。"
    )
    p[28].text = (
        "机器人接近期间持续把 GM461-E1 深度点云和实时激光转换到 base_link，进行地面/离群点过滤、射线清空、障碍膨胀，并在滚动局部图中保留约 0.5—1 s 的短时记忆。局部路径规划可先采用局部 A* 加平台已有跟踪器，或复用已获批的成熟局部控制器；本项目不自行绕过底层安全链路发布 /cmd_vel。目标丢失、相机超时、里程计失效、局部路径被阻或安全输入未知时立即停止并记录原因。"
    )
    table = document.tables[0]
    cell(table, 1, 1, "学习机器人急停、人工接管和低速操作；用 Viewer/官方 SDK 检查 GM461-E1 RGB、Depth、内参、点云、时间戳与对齐；确认 odom→base_link→camera_link、实时激光、平台局部控制/RCS 和停止语义。采集离线 RGB-D，并整理接口与安全检查表。")
    cell(table, 2, 1, "完成 Tkinter、YAML 词表和中文别名；验证 YOLO-World GPU/CPU 行为及 YOLOv8n ONNX 降级；完成 8—10 类物体离线检测、深度中位数和 camera_link/base_link 局部坐标计算。")
    cell(table, 3, 1, "实现当前视野连续三帧确认；打通 camera_link→base_link→odom；建立带射线清空、障碍膨胀和 0.5—1 s 记忆的滚动局部代价图；实现停靠点、局部短路径和每 0.3—0.5 m 重规划的离线回放。")
    cell(table, 4, 1, "连接文字输入、目标确认、局部定位、滚动建图、局部规划和安全仲裁；测试无目标、目标丢失、局部路径被阻、相机超时、里程计/TF 失效及后端降级。全部现场条件通过后再进行短距离低速演示，并归档配置、日志、截图、录像和测试记录。")
    document.core_properties.subject = "目标可见条件下的短距离无地图局部导航与静态避障"
    document.save(PLAN)


def update_path() -> None:
    document = Document(PATH)
    replace_everywhere(document, {
        "ROS 2": "ROS 1 Noetic",
        "RViz2": "RViz",
        "rosbag2": "rosbag",
        "/text_command": "/object_nav/text_command",
    })
    p = document.paragraphs
    p[3].text = "GM461-E1 目标可见短距离局部导航与静态避障"
    p[4].text = "从文字指令、可见目标确认到滚动局部规划和安全停止的一个月完整技术路径"
    p[7].text = (
        "输入“找一下背包”\n"
        "  → command_parser：backpack + policy\n"
        "  → YOLO-World：当前视野候选框 + confidence\n"
        "  → depth fusion：检测框中央有效深度中位数\n"
        "  → TF：camera_link → base_link → odom\n"
        "  → target confirmer：连续三帧类别、深度和位置稳定\n"
        "  → goal builder：目标前约 1 m 的局部停靠点\n"
        "  → rolling local map：深度/激光、射线清空、膨胀与短时记忆\n"
        "  → local planner/controller：执行 0.3—0.5 m 后重新感知和规划\n"
        "  → safety arbiter：到达，或因无目标/无局部路径/数据失效安全停止"
    )
    p[20].text = "阶段 5：坐标变换与短时局部目标坐标"
    p[21].text = "阶段 6：短时里程计与滚动局部参考系"
    p[22].text = "6.1 无全局地图时仍保留哪些空间信息"
    p[23].text = "阶段 7：当前视野目标确认"
    p[24].text = "7.1 LOCAL_SEARCH 的有限边界"
    p[25].text = (
        "local_search:\n"
        "  target_must_be_visible_before_motion: true\n"
        "  confirmation_frames: 3\n"
        "  max_frames: 6\n"
        "  autonomous_exploration: false\n"
        "  map_waypoints: []\n"
        "# 当前视野未确认目标即失败；未经授权不原地旋转或向前探索。"
    )
    p[26].text = "阶段 8：计算局部停靠点并滚动短距离重规划"
    p[27].text = "8.1 LocalNavigationAdapter 统一接口"
    p[28].text = "阶段 9：连续深度/激光生成滚动局部代价地图"
    p[31].text = "10.1 局部无地图任务状态机"
    p[33].text = "12. ROS 1 Noetic 与 RCS 接口映射"
    p[34].text = "13. 无地图局部导航集中配置建议"
    p[35].text = (
        "project_scope:\n"
        "  global_map_required: false\n"
        "  target_must_be_visible_before_motion: true\n"
        "  maximum_demo_range_m: 8.0\n"
        "detector:\n"
        "  primary_backend: yolo_world\n"
        "  fallback_backend: yolov8n_onnx_cpu\n"
        "  confirmation_frames: 3\n"
        "local_navigation:\n"
        "  frame_id: odom\n"
        "  standoff_m: 1.0\n"
        "  replan_distance_m: 0.4\n"
        "  stop_when_target_lost: true\n"
        "local_map:\n"
        "  size_m: [8.0, 8.0]\n"
        "  resolution_m: 0.05\n"
        "  persistence_s: 0.8\n"
        "  inflation_radius_m: 0.5\n"
        "# 以上为离线起点；真实数值必须由机器人尺寸、制动和现场数据验证。"
    )
    p[43].text = "ROS 1 TF 与里程计：http://wiki.ros.org/tf2_ros ｜ http://wiki.ros.org/navigation/Tutorials/RobotSetup/Odom"
    p[44].text = "ROS 1 costmap_2d 障碍层与滚动窗口：http://wiki.ros.org/costmap_2d"

    tables = document.tables
    cell(tables[0], 0, 1, "个人实习项目｜目标可见条件下的短距离无地图局部导航")
    cell(tables[0], 4, 1, "约 15–25 种配置常见物体；仅已验证且当前可见目标可接近")
    cell(tables[0], 5, 1, "离线滚动局部闭环；低速实机为条件性交付")
    cell(tables[1], 0, 0, "一句话理解：YOLO 回答“当前视野里目标在哪里”，深度和局部 TF 回答“目标相对机器人多远”，里程计提供短时运动估计，深度/激光持续生成滚动局部代价地图，局部规划器每执行一小段便重新感知和规划；系统不依赖预建全局地图，安全系统始终拥有最高否决权。")

    flow = [
        ("0", "接口确认", "相机、短时坐标、安全和局部控制权限是否可用？", "能力清单、odom/TF、降级开关"),
        ("1", "文字输入", "用户要寻找什么？", "结构化 TextCommand"),
        ("2", "RGB-D 采集", "彩色、深度和内参是否同步有效？", "Image、Depth、CameraInfo"),
        ("3", "目标检测", "目标是否已在当前视野中？", "候选框、置信度"),
        ("4", "深度融合", "目标在相机前方什么三维位置？", "Point3D_camera"),
        ("5", "局部坐标", "目标相对机器人/短时里程计在哪里？", "PoseStamped(odom)"),
        ("6", "连续确认", "候选是否连续三帧稳定？", "ConfirmedTarget/失败"),
        ("7", "滚动局部图", "局部窗口哪里可通行？", "local_costmap"),
        ("8", "局部规划", "到停靠点的下一小段如何走？", "局部路径/控制任务"),
        ("9", "滚动重规划", "目标和障碍变化后如何修正？", "更新目标、地图和短路径"),
        ("10", "安全执行", "建议动作是否允许？", "NORMAL/SLOW/STOP/FAILED"),
        ("11", "反馈验收", "结果能否解释和复现？", "日志、截图、录像、报告"),
    ]
    for row, values in enumerate(flow, 1):
        for column, value in enumerate(values):
            cell(tables[2], row, column, value)
    cell(tables[3], 9, 1, "ROS PointCloud2/LaserScan + costmap_2d，或 Python 局部栅格")
    cell(tables[3], 9, 2, "以机器人为中心的滚动局部代价表示")
    cell(tables[3], 9, 3, "8 m × 8 m 起步，持续重建和衰减")
    cell(tables[3], 10, 1, "局部 A* + 已有跟踪器，或平台成熟局部控制器/RCS")
    cell(tables[3], 10, 2, "局部短路径、分段执行和高频重规划")
    cell(tables[3], 10, 3, "不自行绕过安全层发布 /cmd_vel")
    cell(tables[3], 11, 1, "rosbag/CSV/JSON/RViz")

    stage0 = tables[5]
    cell(stage0, 2, 1, "GM461-E1 实物；机器人参数表；Viewer/SDK；odom/TF、实时激光、平台局部控制或 RCS 文档；急停、触边、抱闸和接管方式。")
    cell(stage0, 3, 1, "枚举并记录 RGB/Depth/内参/点云；确认深度单位和同步；绘制 odom→base_link→camera_link；确认里程计新鲜度、局部目标/停止权限、速度限制和官方恢复流程。全局地图不是输入。")
    cell(stage0, 5, 1, "Percipio Viewer、官方示例、RViz/TF 工具、厂商 RCS 调试工具。")
    cell(tables[6], 3, 1, "odom、base_link、camera_link；里程计更新时间；实时激光；局部目标、取消和停止接口；是否由成熟控制器执行")
    cell(tables[9], 4, 0, "目标确认")
    cell(tables[9], 4, 1, "当前视野帧数、连续确认数、失败原因")
    cell(tables[9], 5, 1, "odom 位姿、局部目标、局部路径、重规划次数和任务状态")

    stage5 = tables[15]
    stage5_values = [
        "把相机目标转换为短时局部导航可用的 odom 坐标，不建立 map 坐标。",
        "TargetPointCamera；camera_link→base_link 安装外参；odom→base_link 短时变换；统一时间戳。",
        "按目标时间戳查询 TF，计算 P_odom = T_odom_base × T_base_camera × P_camera；检查方向、高度和数据时效；只在短距离闭环内使用。",
        "TargetPoseLocal / PoseStamped(frame_id=odom) 和 transform_quality。",
        "ROS 1 Noetic tf2 或等价矩阵运算。",
        "geometry_msgs/PoseStamped；无 ROS 时保留 frame_id、stamp、position、orientation。",
        "前/左/右静态目标在 base_link/odom 中方向一致；短距离移动后视觉重定位与里程计变化一致。",
        "TF 缺失、过期或里程计失效时 STOP/FAILED；不沿过期目标继续盲行。",
    ]
    for row, value in enumerate(stage5_values, 1): cell(stage5, row, 1, value)
    cell(tables[16], 0, 0, "最常见错误：camera_link、base_link 和 odom 坐标轴方向混淆会让目标左右或前后颠倒。用前/左/右三个已知目标验证，并在短距离移动后复核时间戳和 TF。")

    stage6 = tables[17]
    stage6_values = [
        "在没有全局地图时提供短时运动估计和滚动局部图参考系。",
        "底盘 Odometry、odom→base_link、当前速度、时间戳和健康状态。",
        "读取里程计并检查新鲜度；局部代价地图以 base_link 为滚动中心，目标可在 odom 中短时保持；持续视觉重定位优先于长期里程计积分。",
        "RobotPoseOdom、速度、里程计质量、局部窗口原点。",
        "nav_msgs/Odometry、tf2、平台底盘/RCS；不加载 map_server 或 AMCL。",
        "/odom、/tf、/scan 或厂商等价接口。",
        "短距离前进/转动方向正确；里程计超时可识别；局部图始终随机器人滚动。",
        "里程计不可用时只做静态检测和离线回放，不执行局部运动。",
    ]
    for row, value in enumerate(stage6_values, 1): cell(stage6, row, 1, value)
    t18 = tables[18]
    for col, value in enumerate(("空间信息", "更新方式", "解决的问题", "项目决策")): cell(t18, 0, col, value)
    for col, value in enumerate(("预建全局地图", "不加载、不维护", "长距离连通性与跨房间搜索", "不依赖；超出本项目边界")): cell(t18, 1, col, value)
    for col, value in enumerate(("滚动局部代价地图", "深度/激光持续更新并随机器人滚动", "当前窗口内短距离路径与静态避障", "主线实现或接入")): cell(t18, 2, col, value)

    stage7 = tables[19]
    stage7_values = [
        "在任何运动前确认目标已在当前视野内稳定存在。",
        "规范目标、navigation_mode、连续 RGB-D 检测、局部目标点和安全状态。",
        "在固定数量的当前视野帧中更新同类候选；连续三帧类别、置信度、深度和局部位置稳定才确认；无目标或确认中断则失败，不产生探索运动。",
        "confirmation_count、ConfirmedTarget 或 TARGET_NOT_VISIBLE。",
        "有限状态机和位置稳定性门槛；不配置地图观察点。",
        "TargetConfirmer 不调用 NavigationAdapter。",
        "单帧误检不确认；连续三帧才通过；当前视野无目标时零导航调用并进入 FAILED。",
        "若获批原地扫描，可作为独立后续能力；首版默认由人工调整视角或使用离线多视角回放。",
    ]
    for row, value in enumerate(stage7_values, 1): cell(stage7, row, 1, value)

    stage8 = tables[20]
    stage8_values = [
        "在局部窗口中到达目标附近的停靠距离，并持续重定位和重规划。",
        "目标 odom 相对坐标、机器人短时位姿、滚动局部代价图、机器人外形和停靠距离。",
        "生成朝向目标且相距约 1 m 的候选停靠点；在膨胀后的局部图中规划短路径；每次只执行约 0.3—0.5 m，随后重新检测目标、更新局部图和路径。",
        "LocalGoal、局部路径、控制 task_id、重规划和停止原因。",
        "局部 A* + 平台已有跟踪器，或经批准的成熟局部控制器/RCS。",
        "get_pose/navigate_to_pose/get_navigation_status/request_stop；目标 frame 为 odom 或平台明确支持的局部 frame。",
        "路径不穿越膨胀障碍；目标持续可见；分段执行可停止；到达配置停靠距离。",
        "局部窗口内无路径、目标丢失或目标移出有效范围立即 STOP/FAILED，不猜测窗口外绕行路线。",
    ]
    for row, value in enumerate(stage8_values, 1): cell(stage8, row, 1, value)
    cell(tables[22], 1, 1, "作为短距离导航的唯一环境表示持续更新当前窗口，而不是给全局路径做补充。")
    cell(tables[22], 5, 1, "ROS 1 costmap_2d ObstacleLayer/VoxelLayer，或 Python/NumPy 简化局部栅格；Open3D 用于离线验证。")
    cell(tables[24], 1, 1, "在每次局部路径执行前检查平台安全、数据新鲜度、里程计、目标可见性和局部路径健康。")
    cell(tables[24], 2, 1, "局部路径、目标可见性、最近障碍距离、里程计/相机状态、急停/触边/抱闸、接口健康和人工授权。")
    cell(tables[24], 3, 1, "固定优先级：急停/硬件安全 > 数据或里程计失效 > 目标丢失 > 无局部路径 > 近障减速/停止 > 正常；每小段执行前重新仲裁。")

    states = tables[25]
    if not any(row.cells[0].text == "LOCAL_PLAN" for row in states.rows):
        new_row = states.add_row()
        states.rows[5]._tr.addprevious(new_row._tr)
    cell(states, 2, 3, "可检测且策略允许→LOCAL_SEARCH；detect_only→只展示；否则 FAILED")
    cell(states, 3, 0, "LOCAL_SEARCH")
    cell(states, 3, 1, "当前视野目标未确认")
    cell(states, 3, 2, "固定帧检测和连续确认；不产生运动")
    cell(states, 3, 3, "三帧确认→TARGET_CONFIRMED；无目标/中断→FAILED")
    cell(states, 4, 2, "生成 odom 局部目标并进入局部规划")
    cell(states, 4, 3, "有效→LOCAL_PLAN；失效→FAILED")
    cell(states, 5, 0, "LOCAL_PLAN")
    cell(states, 5, 1, "目标、里程计和滚动局部图有效")
    cell(states, 5, 2, "生成停靠点和下一段局部短路径")
    cell(states, 5, 3, "有路径→APPROACH；无解→STOP/FAILED")
    cell(states, 6, 0, "APPROACH")
    cell(states, 6, 1, "局部控制任务已接受")
    cell(states, 6, 2, "分段执行、持续目标重定位、局部图更新和重规划")
    cell(states, 6, 3, "到达→ARRIVED；目标/里程计/接口失效或路径被阻→STOP/FAILED")
    cell(tables[26], 2, 1, "原始 RGB-D/点云/激光、指令、检测、局部目标、里程计、局部图、局部路径、重规划、控制状态、安全决策和反馈。")
    cell(tables[26], 3, 1, "统一时间戳；记录 task_id、阶段、原因码和配置版本；保存关键图像、滚动局部图和路径；用 rosbag/CSV/JSON 回放。")
    cell(tables[26], 5, 1, "rosbag、CSV/JSON、RViz、Python 测试工具。")

    mapping = tables[27]
    cell(mapping, 0, 1, "ROS 1 Noetic 优先接口")
    cell(mapping, 3, 1, "PointStamped/PoseStamped(odom)")
    cell(mapping, 4, 1, "平台局部目标 action/service 或 odom-frame goal")
    cell(mapping, 5, 1, "costmap_2d、PointCloud2、LaserScan")

    weeks = tables[28]
    cell(weeks, 1, 1, "D1 安全与设备；D2 Viewer；D3 SDK/ROS 1；D4 RGB-D/内参/外参；D5 odom、TF、激光、局部控制/RCS 和离线数据归档。")
    cell(weeks, 3, 1, "D11 TF 到 odom；D12 当前视野三帧确认；D13 滚动局部栅格与障碍膨胀；D14 局部停靠点与短路径；D15 分段执行和滚动重规划测试。")
    cell(weeks, 3, 2, "无目标不运动；三帧确认和物体策略生效；局部目标、滚动图、短路径、目标丢失和阻塞停止行为正确。")

    acceptance = tables[29]
    cell(acceptance, 6, 2, "可见目标边界")
    cell(acceptance, 6, 3, "不访问观察点、不自主探索，零导航调用并进入 TARGET_NOT_VISIBLE/FAILED。")
    cell(acceptance, 7, 3, "单帧保持 LOCAL_SEARCH；三帧满足质量和策略后才确认。")
    cell(acceptance, 8, 3, "不生成 odom 局部目标，失败停止。")
    cell(acceptance, 11, 3, "在滚动局部图中重新规划；无安全局部路径时停止并报告 LOCAL_PATH_BLOCKED。")
    cell(acceptance, 12, 1, "相机超时/里程计丢失")
    cell(acceptance, 13, 3, "在约 1 m 配置距离内停止，并保存目标、里程计、局部图、路径和录像。")
    cell(tables[30], 4, 0, "目标在 odom 中短时漂移")
    cell(tables[30], 4, 1, "时间戳、TF、轮速里程计和持续视觉更新")
    cell(tables[30], 4, 2, "用过期 TF，或里程计短时误差过大")
    cell(tables[30], 4, 3, "按目标 stamp 查询；每小段后视觉重定位；超时即停")
    cell(tables[31], 0, 0, "最终边界：本项目是目标可见条件下的短距离无地图局部导航 Demo，不构成机器人安全认证，不具备跨房间搜索、未知环境探索或绕长墙能力。当前视野无目标、目标丢失、局部窗口无安全路径、数据过期或安全状态未知时必须停止。视觉结果只提供目标和局部障碍建议，不替代急停、触边、抱闸、激光停障和人工接管。")
    document.core_properties.subject = "目标可见条件下的短距离无地图局部导航与静态避障"
    document.save(PATH)


def main() -> None:
    update_plan()
    update_path()
    print(f"updated: {PLAN}")
    print(f"updated: {PATH}")


if __name__ == "__main__":
    main()
